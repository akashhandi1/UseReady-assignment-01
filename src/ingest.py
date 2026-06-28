"""Document ingestion & file-type routing.

Two clean buckets (verified during EDA):
  * .docx  -> real, extractable text (incl. tables); zero embedded images.
  * .png   -> scanned image; read directly by the multimodal model.

OCR (pytesseract) is provided only as an offline fallback for the image path.
"""
from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Optional


@dataclass
class DocumentInput:
    """Normalized representation of one input document."""

    file_name: str          # stem, no extension (matches CSV "File Name")
    modality: str           # "text" or "image"
    text: Optional[str] = None
    image_bytes: Optional[bytes] = None
    mime_type: Optional[str] = None


def _extract_docx_text(path: Path) -> str:
    """Extract paragraph + table text from a .docx in reading order."""
    from docx import Document

    doc = Document(str(path))
    parts: list[str] = []

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts).strip()


def load_document(path: str | Path) -> DocumentInput:
    """Load a single document, routing by file extension."""
    path = Path(path)
    stem = path.stem
    ext = path.suffix.lower()

    # ".pdf.docx" -> stem keeps ".pdf"; strip it so the name matches the CSV.
    if stem.lower().endswith(".pdf"):
        stem = stem[: -len(".pdf")]

    if ext == ".docx":
        return DocumentInput(file_name=stem, modality="text", text=_extract_docx_text(path))

    if ext in {".png", ".jpg", ".jpeg"}:
        mime = "image/png" if ext == ".png" else "image/jpeg"
        return DocumentInput(
            file_name=stem,
            modality="image",
            image_bytes=path.read_bytes(),
            mime_type=mime,
        )

    raise ValueError(f"Unsupported file type: {path.name}")


def ocr_image(image_bytes: bytes) -> str:
    """Offline OCR fallback for images (used only when no vision model)."""
    import io

    import pytesseract
    from PIL import Image

    return pytesseract.image_to_string(Image.open(io.BytesIO(image_bytes)))


SUPPORTED_EXTS = {".docx", ".png", ".jpg", ".jpeg"}


def list_documents(folder: str | Path) -> list[Path]:
    """Return all supported document paths in a folder, sorted by name."""
    folder = Path(folder)
    return sorted(p for p in folder.iterdir() if p.suffix.lower() in SUPPORTED_EXTS)
